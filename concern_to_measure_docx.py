#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Concern to Measure -- Word (.docx) worksheet generator.
Uses python-docx to produce an A4 printable handout.
"""

from docx import Document
from docx.shared import Pt, Mm, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colours ──────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1a, 0x2e, 0x4a)
BLUE   = RGBColor(0x2d, 0x6a, 0x9f)
LGREY  = RGBColor(0xe8, 0xf0, 0xf7)
MID    = RGBColor(0xc5, 0xd8, 0xec)
BORDER = RGBColor(0xa0, 0xb8, 0xd0)
CAPTION= RGBColor(0x55, 0x55, 0x55)
WHITE  = RGBColor(0xff, 0xff, 0xff)


# ── XML helpers ───────────────────────────────────────────────────────────────

def rgb_hex(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def set_cell_bg(cell, color: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  rgb_hex(color))
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)]:
        if val is not None:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),   val.get("val",   "single"))
            el.set(qn("w:sz"),    val.get("sz",    "4"))
            el.set(qn("w:color"), val.get("color", "auto"))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_table_no_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr")) or OxmlElement("w:tblPr")
    bdr   = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        bdr.append(el)
    tblPr.append(bdr)


def para_space(para, before_pt=0, after_pt=0):
    pPr = para._p.get_or_add_pPr()
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:before"), str(int(before_pt * 20)))
    spc.set(qn("w:after"),  str(int(after_pt  * 20)))
    pPr.append(spc)


def set_col_widths(table, widths_mm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW  = OxmlElement("w:tcW")
            tcW.set(qn("w:w"),    str(int(Mm(widths_mm[i]).twips)))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)


# ── Style helpers ─────────────────────────────────────────────────────────────

def add_heading(doc, text):
    """Dark navy full-width banner."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_no_borders(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, NAVY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = WHITE
    cell.paragraphs[0]._p.get_or_add_pPr()
    para_space(p, before_pt=2, after_pt=2)
    return tbl


def add_label(doc, text, note=None):
    """Bold navy label paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = NAVY
    para_space(p, before_pt=6, after_pt=0)
    if note:
        pn = doc.add_paragraph(note)
        pn.runs[0].italic = True
        pn.runs[0].font.size = Pt(7.5)
        pn.runs[0].font.color.rgb = CAPTION
        para_space(pn, before_pt=0, after_pt=0)
    return p


def add_write_line(doc, n=1):
    """Ruled writing line(s) with generous vertical space."""
    for _ in range(n):
        p = doc.add_paragraph()
        para_space(p, before_pt=14, after_pt=0)
        # Bottom border = the ruled line
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "4")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), rgb_hex(BORDER))
        pBdr.append(bot)
        pPr.append(pBdr)


def add_check(doc, label, note=None):
    p = doc.add_paragraph()
    run = p.add_run("\u2610  " + label)
    run.font.size = Pt(8.5)
    para_space(p, before_pt=1, after_pt=1)
    if note:
        pn = doc.add_paragraph("      " + note)
        pn.runs[0].italic = True
        pn.runs[0].font.size = Pt(7.5)
        pn.runs[0].font.color.rgb = CAPTION
        para_space(pn, before_pt=0, after_pt=0)


def add_body_bold(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = NAVY
    para_space(p, before_pt=6, after_pt=2)


def add_note(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(7.5)
    p.runs[0].font.color.rgb = CAPTION
    para_space(p, before_pt=0, after_pt=2)


def add_hr(doc, color=BORDER):
    p = doc.add_paragraph()
    para_space(p, before_pt=2, after_pt=2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), rgb_hex(color))
    pBdr.append(bot)
    pPr.append(pBdr)


# ── Process flow table ────────────────────────────────────────────────────────

def add_process_flow(doc):
    steps = ["A\nRaw statement\n(verbatim)",
             "B\nClassify\nconcern type",
             "C\nReframe as\na concern",
             "D\nScenario\n(env/stim/resp)",
             "E\nMeasure\n(metric+evidence)",
             "F\nRequirement\ninstance (SMART)"]
    tbl = doc.add_table(rows=1, cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_no_borders(tbl)
    for i, (cell, step) in enumerate(zip(tbl.rows[0].cells, steps)):
        parts = step.split("\n")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(parts[0])
        r1.bold = True
        r1.font.size = Pt(7.5)
        if len(parts) > 1:
            p.add_run("\n" + "\n".join(parts[1:])).font.size = Pt(7)
        set_cell_bg(cell, MID if i in (0, 5) else LGREY)
        para_space(p, before_pt=4, after_pt=4)
        # thin border around each cell
        bdr_val = {"val": "single", "sz": "4", "color": rgb_hex(BORDER)}
        set_cell_borders(cell, top=bdr_val, bottom=bdr_val,
                         left=bdr_val, right=bdr_val)
    return tbl


# ── Scenario table ────────────────────────────────────────────────────────────

def add_scenario_table(doc):
    rows_data = [
        ("Environment (assumptions):",
         "Load / degraded mode / location / user group"),
        ("Stimulus (event / load / failure / threat / change):",
         "What triggers the scenario?"),
        ("Scope (service / transaction / user group / component):",
         "Boundary of concern"),
        ("Response (observable behaviour):",
         "What the system does in response"),
        ("Response measure (how success is measured):",
         "Quantitative target \u2014 feeds directly into the Requirement instance"),
    ]
    # 2 cols per row: label | hint | write line spans full width in next row
    tbl = doc.add_table(rows=len(rows_data) * 2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_no_borders(tbl)
    W = 173  # total usable mm on A4 with 18mm margins each side
    set_col_widths(tbl, [W * 0.58, W * 0.42])

    for i, (lbl, hint) in enumerate(rows_data):
        # Row A: label + hint
        cell_lbl  = tbl.cell(i * 2, 0)
        cell_hint = tbl.cell(i * 2, 1)
        pl = cell_lbl.paragraphs[0]
        rl = pl.add_run(lbl)
        rl.bold = True
        rl.font.size = Pt(8.5)
        rl.font.color.rgb = NAVY
        para_space(pl, before_pt=4, after_pt=0)

        ph = cell_hint.paragraphs[0]
        rh = ph.add_run(hint)
        rh.italic = True
        rh.font.size = Pt(7.5)
        rh.font.color.rgb = CAPTION
        para_space(ph, before_pt=4, after_pt=0)

        # Row B: write line spanning both columns
        merged = tbl.cell(i * 2 + 1, 0).merge(tbl.cell(i * 2 + 1, 1))
        pw = merged.paragraphs[0]
        para_space(pw, before_pt=14, after_pt=2)
        pPr = pw._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "4")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), rgb_hex(BORDER))
        pBdr.append(bot)
        pPr.append(pBdr)

    return tbl


# ── Measure table ─────────────────────────────────────────────────────────────

def write_line_in_cell(cell, spaceBefore=14):
    p = cell.paragraphs[0]
    para_space(p, before_pt=spaceBefore, after_pt=2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), rgb_hex(BORDER))
    pBdr.append(bot)
    pPr.append(pBdr)


def add_measure_table(doc):
    W = 173
    tbl = doc.add_table(rows=8, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_no_borders(tbl)
    set_col_widths(tbl, [W * 0.60, W * 0.40])

    def label_cell(cell, text, bold=True):
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(8.5)
        r.font.color.rgb = NAVY
        para_space(p, before_pt=6, after_pt=0)

    # Row 0: shall statement label (merged)
    merged0 = tbl.cell(0, 0).merge(tbl.cell(0, 1))
    label_cell(merged0, 'Requirement instance (SMART \u201cshall\u201d statement):')

    # Row 1: write line (merged)
    merged1 = tbl.cell(1, 0).merge(tbl.cell(1, 1))
    write_line_in_cell(merged1)

    # Row 2: write line 2 (merged)
    merged2 = tbl.cell(2, 0).merge(tbl.cell(2, 1))
    write_line_in_cell(merged2)

    # Row 3: Metric label | Measurement point label
    label_cell(tbl.cell(3, 0), "Metric(s) & target(s):")
    label_cell(tbl.cell(3, 1), "Measurement point (where measured):")

    # Row 4: write lines in both cells
    write_line_in_cell(tbl.cell(4, 0))
    write_line_in_cell(tbl.cell(4, 1))

    # Row 5: Time window | Owner
    label_cell(tbl.cell(5, 0), "Time window / conditions:")
    label_cell(tbl.cell(5, 1), "Owner: _________________  Due / gate: _____________")

    # Row 6: write lines
    write_line_in_cell(tbl.cell(6, 0))
    write_line_in_cell(tbl.cell(6, 1))

    # Row 7: evidence (merged, shaded)
    merged7 = tbl.cell(7, 0).merge(tbl.cell(7, 1))
    set_cell_bg(merged7, LGREY)
    p = merged7.paragraphs[0]
    evidence = ("\u2610 Test   \u2610 Monitoring   \u2610 Audit   "
                "\u2610 Drill   \u2610 Other: __________")
    r = p.add_run("Evidence method:  ")
    r.bold = True
    r.font.size = Pt(8.5)
    p.add_run(evidence).font.size = Pt(8.5)
    para_space(p, before_pt=5, after_pt=5)

    return tbl


# ── Header row ────────────────────────────────────────────────────────────────

def add_header_row(doc):
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_no_borders(tbl)
    W = 173
    set_col_widths(tbl, [W / 3, W / 3, W / 3])
    labels = [
        "Who / role: ___________________________",
        "Date: ____________________",
        "System / service: ___________________________",
    ]
    for cell, lbl in zip(tbl.rows[0].cells, labels):
        p = cell.paragraphs[0]
        r = p.add_run(lbl)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = NAVY
        para_space(p, before_pt=2, after_pt=2)
    return tbl


# ── Main ──────────────────────────────────────────────────────────────────────

def build_docx(out_path="concern_to_measure.docx"):
    doc = Document()

    # Page size A4 with margins
    section = doc.sections[0]
    section.page_width  = Mm(210)
    section.page_height = Mm(297)
    section.left_margin   = Mm(18)
    section.right_margin  = Mm(18)
    section.top_margin    = Mm(18)
    section.bottom_margin = Mm(14)

    # Default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    rt = title.add_run("Concern to Measure")
    rt.bold = True
    rt.font.size = Pt(18)
    rt.font.color.rgb = NAVY
    para_space(title, before_pt=0, after_pt=4)

    add_hr(doc, BLUE)

    add_note(doc,
        "Convert stakeholder language into structured concerns, scenarios, and measurable "
        "requirement instances. Follows ISO/IEC/IEEE 42010:2011 and TOGAF / BCS guidance. "
        "Allow 10\u201320 minutes per statement.")

    add_process_flow(doc)

    add_note(doc,
        "Figure: Raw statement \u2192 Classify \u2192 Reframe \u2192 "
        "Scenario \u2192 Measure \u2192 Requirement instance")

    # ── Section 1 ─────────────────────────────────────────────────────────────
    add_heading(doc, "1 \u00b7 Raw Statement")
    doc.add_paragraph()  # small gap
    add_header_row(doc)

    add_label(doc, "Raw statement (verbatim):")
    add_write_line(doc, 2)

    add_label(doc, "Context \u2014 when / where / which user journey or business scenario:",
              note="Include environment assumptions: peak / normal / degraded  \u00b7  "
                   "locations  \u00b7  user groups.")
    add_write_line(doc, 2)

    add_label(doc, "Impact \u2014 what is at risk (harm / cost / delay / compliance / operations):")
    add_write_line(doc, 2)

    # ── Section 2 ─────────────────────────────────────────────────────────────
    add_heading(doc, "2 \u00b7 Classify Concern Type  [tick all that apply]")

    add_body_bold(doc, "Functional attributes \u2014 what the system must do")
    for item in [
        "Workflow / process step (create, approve, fulfil, reconcile)",
        "Data capture / validation / calculations",
        "Reporting / analytics output",
        "Integration / API / interface behaviour",
        "Roles / permissions / approvals (as a functional need)",
        "Notifications / messaging",
        "Search / retrieval",
        "Audit trail as a feature (who did what, when)",
        "Other: _______________________________",
    ]:
        add_check(doc, item)

    add_body_bold(doc,
        "Quality attributes / NFR types \u2014 how well it must do it  "
        "(ISO/IEC 25010 primary labels \u2192 BCS NFR terms)")
    for lbl, note in [
        ("Performance efficiency (ISO 25010) \u2192 Performance (BCS: throughput / response time)", None),
        ("Reliability (ISO 25010) \u2192 Availability / reliability / recoverability (BCS)", None),
        ("Security (ISO 25010) \u2192 Security / privacy / information assurance (BCS)", None),
        ("Usability (ISO 25010) \u2192 Usability / customer experience  (accessibility often sits here)", None),
        ("Compatibility (ISO 25010) \u2192 Interoperability / integratability (BCS)", None),
        ("Maintainability (ISO 25010) \u2192 Maintainability / modifiability / evolvability (BCS)", None),
        ("Portability (ISO 25010) \u2192 Portability (BCS)",
         "Tip: keep type vs instance clear \u2014 e.g. \u2018throughput\u2019 as type vs "
         "\u2018100 tps\u2019 as instance (Avancier)."),
    ]:
        add_check(doc, lbl, note)

    add_body_bold(doc,
        "Planning & governance attributes \u2014 constraints, risks, controls  "
        "(ISO 42010 \u00a74.2)")
    for item in [
        "Cost / affordability",
        "Schedule / deadline / time-to-market",
        "Feasibility / known limitations",
        "Regulatory / compliance constraint",
        "Risk / assumptions / issues / dependencies  (RAID prompts)",
        "Standards / policy / principles constraint",
        "Decision authority required (who must approve / sign off)",
        "Supplier / contract / governance control requirement",
        "Other: _______________________________",
    ]:
        add_check(doc, item)

    add_check(doc, "Possible conflict (forces)?   \u2610 Yes   \u2610 No")
    add_note(doc, "If yes, name competing concerns and treat them as forces "
                  "rather than personal disagreement.")

    # ── Section 3 ─────────────────────────────────────────────────────────────
    add_heading(doc, "3 \u00b7 Reframe as a Concern  [neutral, actionable]")
    add_note(doc,
        "Template: \u201cOur concern is [topic] so that [stakeholder outcome], under [context].\u201d  "
        "\u2014 ISO 42010 defines concern as a stakeholder interest; "
        "TOGAF emphasises concerns as roots of requirements decomposition.")
    add_label(doc, "Reframed concern:")
    add_write_line(doc, 2)

    # ── Section 4 ─────────────────────────────────────────────────────────────
    add_heading(doc, "4 \u00b7 Scenario  [make it refutable \u2014 SEI QAW / ATAM structure]")
    add_scenario_table(doc)

    # ── Section 5 ─────────────────────────────────────────────────────────────
    add_heading(doc, "5 \u00b7 Measure  [requirement instance + evidence]")
    add_note(doc,
        "BCS: NFRs are usually quantitatively measurable; link each SMART \u201cshall\u201d "
        "statement to an acceptance test, monitoring threshold, or audit criterion.")
    add_measure_table(doc)

    # ── Footer note ───────────────────────────────────────────────────────────
    add_hr(doc, BORDER)
    add_note(doc,
        "References: ISO/IEC/IEEE 42010:2011 \u00b7 ISO/IEC 25010:2011 \u00b7 "
        "TOGAF\u00ae (The Open Group) \u00b7 BCS Non-Functional Requirements guide \u00b7 "
        "SEI Quality Attribute Workshop (QAW) / ATAM \u00b7 Avancier Limited NFR guidance")

    doc.save(out_path)
    print("Word file written -> " + out_path)


if __name__ == "__main__":
    build_docx("concern_to_measure.docx")
